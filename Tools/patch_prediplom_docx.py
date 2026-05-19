# -*- coding: utf-8 -*-
"""Patch МАКЕТ Преддипломная практика.docx: задания п.3–4 и формулировки «СОДЕРЖАНИЯ» под игровой проект."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "МАКЕТ Преддипломная практика.docx"
BACKUP = ROOT / "МАКЕТ Преддипломная практика.backup.docx"

LINE3 = (
    "Спроектировать программный комплекс (игровое приложение) «HonorKitchen» "
    "на базе Unreal Engine 5: архитектура модулей, процедурная генерация игрового уровня, "
    "искусственный интеллект противника, игровая логика и пользовательский интерфейс."
)
LINE4 = (
    "Разработать программный комплекс (игровое приложение) «HonorKitchen»: "
    "реализация на C++ и Blueprints в UE 5, сборка под Windows, тестирование, "
    "подготовка дистрибутива и пользовательской документации."
)

# В исходном шаблоне после «систему» — четыре U+2026 и одна точка
DOTS = "\u2026\u2026\u2026\u2026."

REPLACEMENTS: list[tuple[str, str]] = [
    (f"Спроектировать информационную систему {DOTS}", LINE3),
    (f"Разработать информационную систему {DOTS}", LINE4),
    (
        "РАЗДЕЛ II. ПРОЕКТИРОВАНИЕ СИСТЕМЫ (ПРИЛОЖЕНИЯ, САЙТА И Т.П.) КОМПЬЮТЕРНОЙ МОДЕЛИ",
        "РАЗДЕЛ II. ПРОЕКТИРОВАНИЕ ИГРОВОГО ПРИЛОЖЕНИЯ (ПРОГРАММНОГО ПРОДУКТА)",
    ),
    (
        "РАЗДЕЛ III. РАЗРАБОТКА СИСТЕМЫ (ПРИЛОЖЕНИЯ, САЙТА И Т.П.) КОМПЬЮТЕРНОЙ МОДЕЛИ",
        "РАЗДЕЛ III. РАЗРАБОТКА ИГРОВОГО ПРИЛОЖЕНИЯ (ПРОГРАММНОГО ПРОДУКТА)",
    ),
    (
        "2.2. Создание структурной модели будущей системы (приложения, сайта и т.п.)",
        "2.2. Создание структурной модели будущего игрового приложения (программного продукта)",
    ),
    (
        "1.2. Назначения и цели создания системы (приложения, сайта и т.п.)",
        "1.2. Назначение и цели создания игрового приложения (программного продукта)",
    ),
    (
        "3.4. Тестирование приложения на наличие ошибок и совместимость с различными операционными системами",
        "3.4. Тестирование игрового приложения на наличие ошибок и работоспособность на целевой платформе (Windows)",
    ),
    (
        "ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ НА ПРЕДДИПЛОИНУЮ ПРАКТИКУ",
        "ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ НА ПРЕДДИПЛОМНУЮ ПРАКТИКУ",
    ),
]


def _replace_in_container(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        p.text = p.text.replace(old, new)
                        n += 1
    return n


def _replace_in_header_footer(doc: Document, old: str, new: str) -> int:
    n = 0
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                if old in p.text:
                    p.text = p.text.replace(old, new)
                    n += 1
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if old in p.text:
                                p.text = p.text.replace(old, new)
                                n += 1
    return n


def main() -> None:
    if not DOCX.is_file():
        raise SystemExit(f"Нет файла: {DOCX}")

    if not BACKUP.is_file():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    for old, new in REPLACEMENTS:
        c = _replace_in_container(doc, old, new)
        c += _replace_in_header_footer(doc, old, new)
        if c:
            print(f"OK ({c}x): {old[:52]}...")
        elif old.startswith("Спроектировать") and LINE3 in "\n".join(p.text for p in doc.paragraphs):
            print(f"[skip уже обновлено] {old[:40]}...")
        else:
            print(f"[skip не найдено] {old[:52]}...")

    doc.save(DOCX)
    print(f"Сохранено: {DOCX.name}")


if __name__ == "__main__":
    main()
